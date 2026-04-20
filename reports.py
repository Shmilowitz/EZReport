"""
Report generators.  Each class accepts a list of Finding objects and metadata
dict, and writes the formatted report to the given output path.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, FileSystemLoader

from findings import Finding, SEVERITY_ORDER

_TEMPLATES_DIR = Path(__file__).parent / 'templates'


# ── Base ──────────────────────────────────────────────────────────────────────

class BaseReport(ABC):
    def __init__(self, findings: List[Finding], metadata: dict):
        self.findings = sorted(findings, key=lambda f: f.severity_order)
        self.metadata = metadata

    @abstractmethod
    def generate(self, output_path: Path) -> None:
        pass

    def severity_counts(self) -> dict:
        counts = {s: 0 for s in SEVERITY_ORDER}
        for f in self.findings:
            counts[f.severity] = counts.get(f.severity, 0) + 1
        return counts


# ── Markdown ──────────────────────────────────────────────────────────────────

_SEV_ICON = {'critical': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '🔵', 'info': '⚪'}


class MarkdownReport(BaseReport):
    def generate(self, output_path: Path) -> None:
        meta   = self.metadata
        counts = self.severity_counts()
        lines: List[str] = [
            f"# {meta.get('title', 'Penetration Test Report')}",
            '',
            f"**Date:** {meta.get('date', 'N/A')}  ",
            f"**Author:** {meta.get('author', 'N/A')}  ",
            f"**Scope:** {meta.get('scope', 'N/A')}  ",
            '', '---', '', '## Executive Summary', '',
            '| Severity | Count |', '|----------|-------|',
        ]
        for sev in SEVERITY_ORDER:
            lines.append(f'| {sev.capitalize()} | {counts[sev]} |')
        lines += ['', f"**Total findings: {sum(counts.values())}**", '', '---', '']

        for sev in SEVERITY_ORDER:
            group = [f for f in self.findings if f.severity == sev]
            if not group:
                continue
            lines += [f"## {_SEV_ICON.get(sev, '')} {sev.capitalize()} Findings", '']
            for f in group:
                lines += [f'### {f.title}', '',
                           f'**Asset:** `{f.affected_asset}`  ',
                           f'**Severity:** {f.severity.capitalize()}  ']
                if f.cvss_score is not None:
                    lines.append(f'**CVSS Score:** {f.cvss_score}  ')
                if f.cwe:
                    lines.append(f'**CWE:** {f.cwe}  ')
                lines.append(f'**Source:** {f.tool_source}  ')
                if f.evidence:
                    lines += ['', '#### Evidence', '', '```', f.evidence, '```', '']
                if f.remediation:
                    lines += ['#### Remediation', '', f.remediation, '']
                if f.references:
                    lines += ['#### References', '']
                    lines += [f'- {r}' for r in f.references]
                    lines.append('')
                lines += ['---', '']

        output_path.write_text('\n'.join(lines), encoding='utf-8')


# ── HTML ──────────────────────────────────────────────────────────────────────

class HtmlReport(BaseReport):
    def generate(self, output_path: Path) -> None:
        env = Environment(loader=FileSystemLoader(str(_TEMPLATES_DIR)), autoescape=True)
        html = env.get_template('report.html.jinja2').render(
            metadata=self.metadata,
            findings=self.findings,
            severity_counts=self.severity_counts(),
            severity_order=list(SEVERITY_ORDER.keys()),
            total=len(self.findings),
        )
        output_path.write_text(html, encoding='utf-8')


# ── DOCX ──────────────────────────────────────────────────────────────────────

_DOCX_COLORS = {
    'critical': RGBColor(0xDC, 0x35, 0x45),
    'high':     RGBColor(0xFD, 0x7E, 0x14),
    'medium':   RGBColor(0xFF, 0xC1, 0x07),
    'low':      RGBColor(0x19, 0x8C, 0xFF),
    'info':     RGBColor(0x6C, 0x75, 0x7D),
}


def _colored_heading(doc: Document, text: str, color: RGBColor, level: int = 2):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color


class DocxReport(BaseReport):
    def generate(self, output_path: Path) -> None:
        doc  = Document()
        meta = self.metadata

        title_p = doc.add_heading(meta.get('title', 'Penetration Test Report'), 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        t = doc.add_table(rows=3, cols=2)
        t.style = 'Table Grid'
        for i, (label, key) in enumerate([('Date', 'date'), ('Author', 'author'), ('Scope', 'scope')]):
            t.rows[i].cells[0].text = label
            t.rows[i].cells[1].text = meta.get(key, 'N/A')
        doc.add_paragraph()

        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_table(rows=1, cols=2)
        summary.style = 'Table Grid'
        summary.rows[0].cells[0].text = 'Severity'
        summary.rows[0].cells[1].text = 'Count'
        for sev, count in self.severity_counts().items():
            r = summary.add_row()
            r.cells[0].text = sev.capitalize()
            r.cells[1].text = str(count)
        doc.add_paragraph()

        doc.add_heading('Findings', level=1)
        for sev in SEVERITY_ORDER:
            group = [f for f in self.findings if f.severity == sev]
            if not group:
                continue
            color = _DOCX_COLORS.get(sev, RGBColor(0, 0, 0))
            _colored_heading(doc, f'{sev.capitalize()} Findings', color, level=2)
            for f in group:
                _colored_heading(doc, f.title, color, level=3)
                dt = doc.add_table(rows=0, cols=2)
                dt.style = 'Table Grid'
                rows = [('Asset', f.affected_asset), ('Severity', f.severity.capitalize())]
                if f.cvss_score is not None:
                    rows.append(('CVSS Score', str(f.cvss_score)))
                if f.cwe:
                    rows.append(('CWE', f.cwe))
                rows.append(('Source', f.tool_source))
                for label, value in rows:
                    row = dt.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = value
                if f.evidence:
                    doc.add_heading('Evidence', level=4)
                    p = doc.add_paragraph(f.evidence)
                    p.runs[0].font.name = 'Courier New'
                    p.runs[0].font.size = Pt(8)
                if f.remediation:
                    doc.add_heading('Remediation', level=4)
                    doc.add_paragraph(f.remediation)
                if f.references:
                    doc.add_heading('References', level=4)
                    for ref in f.references:
                        doc.add_paragraph(ref, style='List Bullet')
                doc.add_paragraph()

        doc.save(str(output_path))
